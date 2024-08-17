from docx import Document

def export_flashcards(data, filename):
    doc = Document()
    doc.add_heading('Flashcards', 0)

    for index, flashcard in enumerate(data):
        doc.add_heading(f'Flashcard {index + 1}', level=1)
        doc.add_paragraph(f'Question: {flashcard[0]}')  # Assuming first element is the question
        doc.add_paragraph(f'Answer: {flashcard[1]}')    # Assuming second element is the answer
        doc.add_page_break()

    doc.save(filename)