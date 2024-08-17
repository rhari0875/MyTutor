from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def export_quiz(quiz_data, docx_filename):
    document = Document()

    heading_style = document.styles["Heading1"]
    heading_font = heading_style.font
    heading_font.size = Pt(14)  
    body_text_style = document.styles["BodyText"]
    body_text_font = body_text_style.font
    body_text_font.size = Pt(12)  

    document.add_heading("Quiz", level=1)

    for question_data in quiz_data:
        question = question_data["question"]
        possible_answers = question_data["possible_answers"]

        p = document.add_paragraph()
        p.add_run(question).bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        for answer in possible_answers:
            document.add_paragraph(answer, style="ListBullet")

    document.save(docx_filename)